from docx import Document
from docx2pdf import convert
from datetime import datetime
import json
import sys
import os

def replace_text_in_docx(input_path, output_path, replacements):
    doc = Document(input_path)

    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                para.text = para.text.replace(old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in replacements.items():
                    if old in cell.text:
                        cell.text = cell.text.replace(old, new)

    doc.save(output_path)

def convert_to_pdf(docx_path):
    pdf_dir = os.path.dirname(docx_path)
    convert(docx_path, pdf_dir)  # will output a PDF next to the DOCX

# Point d’entrée depuis Laravel
if __name__ == '__main__':
    if len(sys.argv) != 2:
        print("Usage: python RESILIATION.py <json_file>")
        sys.exit(1)

    json_path = sys.argv[1]
    with open(json_path, encoding='utf-8') as f:
        data = json.load(f)

    # Remplacer le texte dans le DOCX
    replace_text_in_docx(
        input_path=data['template_path'],
        output_path=data['output_path'],
        replacements=data['replacements']
    )

    # Convertir en PDF
    convert_to_pdf(data['output_path'])

    print("✅ Document Word + PDF générés :", data['output_path'])
